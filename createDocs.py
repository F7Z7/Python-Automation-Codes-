# I had to create a docx file in a certain format ,which was time consuming
# So i came up with this script to automate the process



import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


source_folder = "E:\\To_Jpg"

#speaker Data
speakers = {
    "Adv. Navaneeth": {
        "data": "Navaneeth Pavithran: Empowering Masses through Law. Navaneeth Pavithran, a beacon of inspiration, is a dynamic social media influencer and passionate justice advocate. With a unique vision, he's making legal knowledge accessible, reshaping our perception of law. Navaneeth's journey is unconventional. Trained as a lawyer, he saw law's limitations in reaching everyday people. To bridge this gap, he harnessed digital media's potential. Through one-minute videos and Instagram posts, he passionately simplifies complex legal jargon for everyone. His motto is simple yet profound: bring law to the common man. Legal literacy, he believes, empowers and uplifts individuals and communities. Discover the man, his mission, the stories fueling his passion, and his groundbreaking impact.",
        "image": os.path.join(source_folder, "nav3.jpg")
    },

    "Siddique Kappan": {
        "data": "In the ever-evolving landscape of journalism, where the quest for truth often navigates through the labyrinth of paradoxes, one name stands out as a beacon of integrity and unwavering commitment – Siddique Kappan. A master storyteller, a guardian of the Fourth Estate, and a voice for the voiceless, Siddique's journey through the realm of journalism has been nothing short of extraordinary. Picture this: A journalist who doesn't merely report the news, but delves deep into the crux of stories, uncovering hidden layers of complexity that boggle the mind. Siddique's words have a way of unraveling paradoxes that surround us, exposing the heart of the matter with a rare blend of compassion and diligence.",
        "image": os.path.join(source_folder, "sid2.jpg")

    },

    "Adhila Nasarin": {
        "data": "Adhila Nasarin, a fearless 23-year-old from Aluva, embodies the unwavering power of love and determination. Her courageous journey defies all odds, driven by a profound desire to reunite with her beloved. Adhila's story mirrors the challenges that many LGBTQ+ individuals confront in their pursuit of acceptance and the freedom to love openly. Her efforts, involving interactions with the police and court, underscore the ongoing struggles of the LGBTQ+ community in their search for love and happiness. Adhila's resilience shines as a beacon of hope, showcasing the enduring strength of the human spirit amid adversity.",
        "image": os.path.join(source_folder, "adhila3.jpg")

    },

    "Anoop Ambika": {
        "data":"Anoop P Ambika is the CEO of Kerala Startup Mission. With over 26 years of rich industry experience, Anoop's proficiency spans domains including international sales and marketing, policy making, AI/ML, Life Sciences, and entrepreneurship. Prior to becoming the CEO of KSUM, Mr. Anoop was the CEO of the biotech firm Genpro Research. Mr. Anoop has co-founded several technology start-ups and also represents Kerala Knowledge Economy Mission in the Kerala State Planning Board. He has served as the patron of Natana, the arts and cultural forum for IT companies in Kerala, highlighting his commitment to fostering creativity and artistic expression. As the key person at Kerala Startup Mission, Anoop is responsible for leading the development and execution of a decision sciences product and services strategy using AI and advanced analytics.",

        "image": os.path.join(source_folder, "anoop4.jpg")

    },

    "Anaswara Ramesh": {
        "data": "Anaswara Ramesh, a remarkable teenager from Kochi, Kerala, has left an indelible mark in record books and garnered prestigious awards for her outstanding achievements. Her journey exemplifies a fearless pursuit of knowledge and innovation. Anaswara's academic commitment, acknowledged with the CCRT Scholarship, and her technological aptitude, showcased by her WhiteHat Jr Scholarship, showcase her diverse talents. She serves as an inspiration, showcasing that determination and a passion for change can dismantle barriers and create a substantial impact.",
        "image": os.path.join(source_folder, "anaswara3.jpg")

    },

    "Priyesh Gopalakrishnan": {
        "data": "With a career spanning 25 years, Priyesh is an accomplished IT leader in the industry specifically in driving Large Scale Program Management, Portfolio Management, Account Management, Digital Transformation, Thought Leadership, Financial and Operations Management, and IT budgeting. Presently, he is the Delivery Director at Cognizant. He has consistently pushed the boundaries of what's possible in the realm of technology. Mr. Priyesh’s contributions extend far beyond the boardroom. He has a keen interest in Psychology and is a certified Life Coach, Agile Coach, Personal Transformation Coach, NLP Master Practitioner, Enneagram Expert and Transactional Analysis Expert. Get ready to be enlightened by our enthusiastic speaker’s journey at TEDxGECRIT.",
        "image": os.path.join(source_folder, "pri5.jpg")

    }
}
#
# for file in os.listdir(source_folder):
#     for speaker in speakers.values():
#         if file == os.path.basename(speaker["image"]):
#             print("valid file name ")

doc = Document()
doc.add_heading('TEDXGECRIT23 Speakers', level=1)

for name, info in speakers.items():


    heading = doc.add_heading(name, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER


    if os.path.exists(info["image"]):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run()
        img_run.add_picture(info["image"], width=Inches(2.5))


    desc = doc.add_paragraph(info["data"])
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER


    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Inches(0.4)

doc.save("TEDxGECRIT23_Speakers.docx")


